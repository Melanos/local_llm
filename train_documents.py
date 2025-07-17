"""
📄 Add Word Documents to RAG Database
This script extracts text from .docx files and adds them to your RAG database.
"""

import os
import requests
import json
import chromadb
from chromadb.utils import embedding_functions
from datetime import datetime
from docx import Document

def setup_rag_database(db_path="./rag_database"):
    """Initialize ChromaDB for storing documents"""
    print("Setting up RAG database...")
    
    client = chromadb.PersistentClient(path=db_path)
    
    embedding_function = embedding_functions.OllamaEmbeddingFunction(
        url="http://localhost:11434",
        model_name="nomic-embed-text",
    )
    
    try:
        # Try to get existing collection first
        collection = client.get_collection(
            name="documents",
            embedding_function=embedding_function
        )
        print(f"✅ Connected to existing RAG database")
    except Exception:
        # Create new collection if it doesn't exist
        collection = client.create_collection(
            name="documents", 
            embedding_function=embedding_function
        )
        print(f"✅ Created new RAG database")
    
    return collection

def extract_text_from_docx(file_path):
    """Extract text content from a Word document"""
    try:
        doc = Document(file_path)
        text_content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n".join(text_content)
        
    except Exception as e:
        print(f"❌ Error extracting text from {file_path}: {e}")
        return None

def chunk_text(text, chunk_size=1000, overlap=200):
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        
        # Try to end at a sentence boundary
        if end < len(text):
            last_period = chunk.rfind('.')
            last_newline = chunk.rfind('\n')
            boundary = max(last_period, last_newline)
            
            if boundary > start + chunk_size // 2:  # Only if boundary is reasonable
                end = start + boundary + 1
                chunk = text[start:end]
        
        chunks.append(chunk.strip())
        start = end - overlap
        
        if end >= len(text):
            break
            
    return chunks

def add_docx_to_rag(file_path):
    """Add a Word document to the RAG database"""
    print(f"📄 Processing: {file_path}")
    
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return False
    
    if not file_path.lower().endswith('.docx'):
        print(f"❌ Not a .docx file: {file_path}")
        return False
    
    # Extract text
    print("  📝 Extracting text content...")
    text_content = extract_text_from_docx(file_path)
    
    if not text_content:
        print("  ❌ No text content extracted")
        return False
    
    print(f"  ✅ Extracted {len(text_content)} characters")
    
    # Chunk the text
    chunks = chunk_text(text_content)
    print(f"  📊 Split into {len(chunks)} chunks")
    
    # Setup database
    collection = setup_rag_database()
    
    # Add chunks to database
    filename = os.path.basename(file_path)
    
    ids = []
    metadatas = []
    documents = []
    
    for i, chunk in enumerate(chunks):
        chunk_id = f"docx_{filename}_chunk_{i}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        chunk_metadata = {
            "type": "word_document",
            "filename": filename,
            "filepath": file_path,
            "chunk_index": i,
            "total_chunks": len(chunks),
            "added_date": datetime.now().isoformat(),
            "file_size": os.path.getsize(file_path)
        }
        
        ids.append(chunk_id)
        metadatas.append(chunk_metadata)
        documents.append(chunk)
    
    # Add to collection
    print("  💾 Adding to database...")
    collection.add(
        documents=documents,
        metadatas=metadatas,
        ids=ids
    )
    
    print(f"  ✅ Added {len(chunks)} chunks to database")
    
    # Show database stats
    count = collection.count()
    print(f"📊 Database now contains {count} total documents")
    
    return True

def show_database_contents():
    """Show what's in the database"""
    try:
        collection = setup_rag_database()
        count = collection.count()
        
        if count == 0:
            print("📊 Database is empty")
            return
        
        print(f"📊 Database contains {count} documents")
        
        # Get sample of documents to show types
        sample = collection.get(limit=min(count, 20))
        file_types = {}
        
        for metadata in sample['metadatas']:
            doc_type = metadata.get('type', 'unknown')
            filename = metadata.get('filename', 'unknown')
            
            if doc_type not in file_types:
                file_types[doc_type] = []
            if filename not in file_types[doc_type]:
                file_types[doc_type].append(filename)
        
        print("\n📁 Document types in database:")
        for doc_type, files in file_types.items():
            print(f"  {doc_type}: {len(files)} files")
            for file in files[:3]:  # Show first 3 files
                print(f"    - {file}")
            if len(files) > 3:
                print(f"    ... and {len(files) - 3} more")
        
    except Exception as e:
        print(f"❌ Error reading database: {e}")

def add_all_docx_from_directory(docs_dir="./Documents"):
    """Add all .docx files from the Documents directory to RAG database"""
    print(f"� Processing all .docx files from: {docs_dir}")
    
    if not os.path.exists(docs_dir):
        print(f"❌ Documents directory not found: {docs_dir}")
        print("💡 Please create the './Documents' directory and add your .docx files there")
        return 0, 0
    
    # Find all .docx files
    docx_files = []
    for file in os.listdir(docs_dir):
        if file.lower().endswith('.docx') and not file.startswith('~$'):  # Skip temp files
            docx_files.append(os.path.join(docs_dir, file))
    
    if not docx_files:
        print(f"❌ No .docx files found in {docs_dir}")
        return 0, 0
    
    print(f"📄 Found {len(docx_files)} .docx files to process")
    
    # Check for existing documents to avoid duplicates
    collection = setup_rag_database()
    existing_docs = collection.get()
    existing_files = set()
    if existing_docs['metadatas']:
        for metadata in existing_docs['metadatas']:
            if metadata.get('type') == 'word_document':
                existing_files.add(metadata.get('filename', ''))
    
    processed_files = 0
    total_chunks = 0
    
    for docx_path in docx_files:
        filename = os.path.basename(docx_path)
        
        # Skip if already processed
        if filename in existing_files:
            print(f"⏭️  Skipping {filename} (already processed)")
            continue
        
        print(f"\n� Processing: {filename}")
        success = add_docx_to_rag(docx_path)
        if success:
            processed_files += 1
            # Count chunks for this file
            file_docs = collection.get(where={"filename": filename})
            if file_docs and file_docs['metadatas']:
                file_chunks = len([m for m in file_docs['metadatas'] if m.get('filename') == filename])
                total_chunks += file_chunks
                print(f"  ✅ Added {file_chunks} chunks")
    
    print(f"\n📊 Processing complete!")
    print(f"   Files processed: {processed_files}")
    print(f"   Total chunks added: {total_chunks}")
    print(f"   Total documents in database: {collection.count()}")
    
    return processed_files, total_chunks

def main():
    """Main function"""
    print("📄 Word Document to RAG Database")
    print("=" * 40)
    
    # Process all documents in the Documents directory
    processed_files, total_chunks = add_all_docx_from_directory("./Documents")
    
    if processed_files > 0:
        print(f"\n🎉 Successfully processed {processed_files} files ({total_chunks} chunks)!")
        print("\n💬 You can now ask questions like:")
        print("   - 'What information is in my documents?'")
        print("   - 'What are my skills?'")
        print("   - 'What work experience do I have?'")
        print("   - 'Tell me about my education'")
        print("   - 'Summarize the documents you have'")
    else:
        print(f"\n⚠️  No new documents were processed")
        print("💡 Available files in Documents folder:")
        if os.path.exists("Documents"):
            for file in os.listdir("Documents"):
                if file.endswith('.docx'):
                    print(f"   - {file}")
        else:
            print("   (Documents folder not found)")
    
    print("\n" + "=" * 40)
    show_database_contents()
    
    print("\n💡 To chat with your documents:")
    print("   python rag_chatbot.py")

if __name__ == "__main__":
    main()
